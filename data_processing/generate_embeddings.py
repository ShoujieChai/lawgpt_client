import os
import json
import sys
from pathlib import Path

# Add the project root directory to Python path
project_root = str(Path(__file__).parent.parent)
if project_root not in sys.path:
    sys.path.append(project_root)

# Now we can import our modules
from config import (
    DOCUMENTS_DIR,
    EMBEDDINGS_DIR,
    SUPPORTED_DOCUMENT_TYPES,
    CHUNK_SIZE
)
from RAG.embedding import get_embedding, get_embedding_stats
from data_processing.preprocess_docs import preprocess_document, split_into_chunks
from typing import Optional

# def read_text_file(file_path: str) -> str:
#     with open(file_path, 'r', encoding='utf-8') as file:
#         return file.read()

def read_txt_file(file_path: str) -> str:
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()

# def read_pdf_file(file_path: str) -> str:
#     text = ""
#     with open(file_path, 'rb') as file:
#         reader = PdfReader(file)
#         for page in reader.pages:
#             text += page.extract_text()
#     return text

def read_pdf_file(file_path: str) -> str:
    try:
        import PyPDF2
    except ImportError:
        raise ImportError("PyPDF2 is required to read PDF files. Please install it with 'pip install PyPDF2'.")
    text = ""
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            text += page.extract_text() or ""
    return text

# def read_docx_file(file_path: str) -> str:
#     text = ""
#     doc = Document(file_path)
#     for para in doc.paragraphs:
#         text += para.text + "\n"
#     return text

def read_docx_file(file_path: str) -> str:
    try:
        import docx
    except ImportError:
        raise ImportError("python-docx is required to read DOCX files. Please install it with 'pip install python-docx'.")
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

# def read_document(file_path: str) -> str:
#     ext = os.path.splitext(file_path)[1].lower()
#     if ext == '.txt':
#         return read_text_file(file_path)
#     elif ext in ['.pdf']:
#         return read_pdf_file(file_path)
#     elif ext in ['.docx', '.doc']:
#         return read_docx_file(file_path)
#     else :
#         print(f"Unsupported file type: {ext}")
#         return None

def read_document(file_path: str) -> Optional[str]:
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.txt':
        return read_txt_file(file_path)
    elif ext == '.pdf':
        return read_pdf_file(file_path)
    elif ext in ('.doc', '.docx'):
        return read_docx_file(file_path)
    else:
        print(f"Unsupported file type: {ext}")
        return None
    


def generate_embeddings_for_document(file_path: str) -> None:
    """
    Read content from a document, preprocess it, split it into chunks,
    and generate embeddings for each chunk.

    """
    os.makedirs(EMBEDDINGS_DIR, exist_ok=True)
    content = read_document(file_path)
    if not content:
        print(f"Could not read content from {file_path}.")
        return
    processed_content = preprocess_document(content)
    chunks = split_into_chunks(processed_content, chunk_size=CHUNK_SIZE)
    for i, chunk in enumerate(chunks):
        embedding = get_embedding(chunk)
        output_file = os.path.join(EMBEDDINGS_DIR, 
            f"{os.path.splitext(os.path.basename(file_path))[0]}_chunk_{i}.json")
        with open(output_file, 'w') as f:
            json.dump({"content": chunk, "embedding": embedding.tolist()}, f)


# def generate_embeddings_for_document(file_path: str) -> None:
#     """
#     Generate embeddings for a document and save them.
#     """
#     os.makedirs(EMBEDDINGS_DIR, exist_ok=True)
#     content = read_document(file_path)
#     if not content:
#         print(f"Could not read content from {file_path}")
#         return
#     processed_content = preprocess_document(content)
#     chunks = split_into_chunks(processed_content, chunk_size=CHUNK_SIZE)
#       for i, chunk in enumerate(chunks):
#         embedding = get_embedding(chunk)
#         output_file = os.path.join(
#             EMBEDDINGS_DIR,
#             f"{os.path.splitext(os.path.basename(file_path))[0]}_chunk_{i}.json"
#         )
#         with open(output_file, 'w') as f:
#             json.dump({
#                 'content': chunk,
#                 'embedding': embedding.tolist()
#             }, f)

def process_all_documents() -> None:
    if not os.path.exists(DOCUMENTS_DIR):
        print(f"Documents directory {DOCUMENTS_DIR} does not exist!")
        return
    for filename in os.listdir(DOCUMENTS_DIR):
        if filename.endswith(SUPPORTED_DOCUMENT_TYPES):
            file_path = os.path.join(DOCUMENTS_DIR, filename)
            print(f"Processing {file_path}...")
            generate_embeddings_for_document(file_path)
            print(f"Finished processing {file_path}.")

if __name__ == "__main__":
    process_all_documents()








